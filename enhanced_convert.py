#!/usr/bin/env python3
"""
Enhanced DOCX to Markdown Converter - 通用工具
Version: 1.1.0

基于以下最佳实践：
1. 标题检测：直接使用 Word 原生样式名
2. 段落一致性：保守策略保留自然分段
3. 表格处理：统一使用 HTML 格式，支持智能纵向合并检测
4. 一致性检查：多维度验证

适用于所有 DOCX 文件的通用转换工具
"""

import os
import sys
import re
import json
import argparse
import zipfile
from pathlib import Path
from typing import Optional, Dict, Tuple, List, Set
from dataclasses import dataclass, field
from enum import Enum
from lxml import etree
from docx import Document
from docx.oxml.ns import qn
from docx.table import Table as DocxTable

# DOCX XML命名空间
NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'


@dataclass
class ProcessingConfig:
    """处理配置"""
    enable_table: bool = True           # 启用表格转换
    fix_chinese_spacing: bool = True   # 修复中文间距
    remove_page_numbers: bool = True    # 移除页码
    keep_headers_footers: bool = False # 保留页眉页脚
    image_enabled: bool = True          # 启用图片提取
    image_to_base64: bool = False       # 图片内嵌Base64（False=导出到文件）
    
    # 标题样式关键词（用于检测非标准标题样式）
    heading_style_keywords: List[str] = field(default_factory=lambda: [
        'heading', 'title', 'toc', 'cover', '标题', '目录'
    ])
    
    # 页码模式
    page_number_patterns: List[str] = field(default_factory=lambda: [
        r'^\d+$',
        r'^第\s*\d+\s*页$',
        r'^Page\s+\d+$',
        r'^\d+\s*/\s*\d+$',
    ])


class TextPostProcessor:
    """文本后处理器"""
    
    @staticmethod
    def fix_chinese_spacing(text: str) -> str:
        """修复中文与拉丁字符的间距"""
        text = re.sub(r'([\u4e00-\u9fff]) ([a-zA-Z])', r'\1\2', text)
        text = re.sub(r'([a-zA-Z]) ([\u4e00-\u9fff])', r'\1\2', text)
        text = re.sub(r'([\u4e00-\u9fff])([a-zA-Z])', r'\1 \2', text)
        text = re.sub(r'([a-zA-Z])([\u4e00-\u9fff])', r'\1 \2', text)
        return text
    
    @staticmethod
    def clean_nbsp(text: str) -> str:
        """移除不间断空格"""
        return re.sub(r'([\u4e00-\u9fff])\u00a0([\u4e00-\u9fff])', r'\1\2', text)
    
    @staticmethod
    def remove_empty_lines(text: str, max_consecutive: int = 1) -> str:
        """
        移除多余的空行
        
        Args:
            text: 输入文本
            max_consecutive: 允许的最大连续空行数（默认1，即最多1个空行分隔）
        """
        # 将连续的空行（仅包含空白字符的行）替换为指定数量
        # 使用正则匹配连续的空行
        lines = text.split('\n')
        result_lines = []
        empty_count = 0
        
        for line in lines:
            if line.strip() == '':
                empty_count += 1
                if empty_count <= max_consecutive:
                    result_lines.append(line)
            else:
                empty_count = 0
                result_lines.append(line)
        
        return '\n'.join(result_lines)


class HeadingDetector:
    """标题检测器 - 基于 Word 原生样式 + 启发式检测"""
    
    # 标准 Word 标题样式模式
    HEADING_PATTERN = re.compile(r'^Heading\s+(\d+)$', re.IGNORECASE)
    # 中文标题模式
    CHINESE_HEADING_PATTERN = re.compile(r'^标题\s*(\d+)$', re.IGNORECASE)
    
    # 启发式小标题模式
    # 可能是小标题的特征
    SUBHEADING_PATTERNS = [
        re.compile(r'^第[一二三四五六七八九十]+[章节条款]'),  # 第X章、第X节
        re.compile(r'^[（\(【]?[A-Z0-9]+[）\)]?[\s　]+'),  # A. B. 1) 等开头
        re.compile(r'^[\u4e00-\u9fff]{2,6}$'),  # 纯中文2-6字
    ]
    
    # 小标题结尾模式（后面内容是大段落的过渡语）
    SUBHEADING_ENDINGS = [
        '如下：', '如下', '包括：', '包括',
        '主要内容', '主要功能', '主要特点',
        '总体架构', '技术架构', '系统架构',
        '建设内容', '项目内容', '工程内容',
    ]
    
    # 排除模式（不应该是标题的短文本）
    EXCLUDE_PATTERNS = [
        re.compile(r'^[\d\.、\)]+$'),  # 纯数字标号
        re.compile(r'^[\(\)（）【】\[]+$'),  # 纯括号
    ]
    
    @classmethod
    def is_heading_style(cls, style_name: str) -> Tuple[bool, int]:
        """
        检测是否为标题样式及其级别
        
        Returns:
            (is_heading, level) - 是标题返回(True, 级别)，否则返回(False, 0)
        """
        if not style_name:
            return False, 0
        
        # 方法1：标准英文模式 "Heading 1", "Heading 2"
        match = cls.HEADING_PATTERN.match(style_name.strip())
        if match:
            return True, int(match.group(1))
        
        # 方法2：中文模式 "标题 1", "标题 2"
        match = cls.CHINESE_HEADING_PATTERN.match(style_name.strip())
        if match:
            return True, int(match.group(1))
        
        return False, 0
    
    @classmethod
    def is_likely_subheading(cls, text: str, next_text: str = "") -> bool:
        """
        启发式检测：判断文本是否可能是小标题
        
        条件：
        1. 文本长度 2-15 字符
        2. 不是排除模式
        3. 不是数字编号列表项（如"（1）" "1." "1、"开头）
        4. 满足以下任一条件：
           - 以特定模式开头（第X章、第X节）
           - 是纯中文2-6字（不含列表编号）
           - 结尾是过渡语（如"如下："）
           - 下一个文本明显是大段落开头
        """
        if not text or len(text) < 2 or len(text) > 15:
            return False
        
        # 排除数字编号列表项模式（这是列表，不是小标题）
        # 例如： （1） xxx、 1. xxx、 1、 xxx、 1）xxx
        list_number_pattern = re.compile(r'^[\（\(【\[第]?\s*\d+[\.\、\)\）\]]+\s*[\u4e00-\u9fff]')
        if list_number_pattern.match(text):
            return False
        
        # 排除模式
        for pattern in cls.EXCLUDE_PATTERNS:
            if pattern.match(text.strip()):
                return False
        
        # 检查是否是纯中文短文本（可能是小标题）
        if re.match(r'^[\u4e00-\u9fff]+$', text):
            # 2-6个汉字，可能是标题
            if 2 <= len(text) <= 6:
                return True
        
        # 检查开头模式（第X章、第X节）
        for pattern in cls.SUBHEADING_PATTERNS:
            if pattern.match(text):
                return True
        
        # 检查结尾过渡语
        for ending in cls.SUBHEADING_ENDINGS:
            if text.endswith(ending):
                return True
        
        # 检查下一个文本是否明显是大段落（正文开头）
        if next_text and len(next_text) > 50:
            # 下一个文本很长，可能是正文
            # 如果当前文本是短文本，可能是标题
            if len(text) <= 15:
                return True
        
        return False
    
    @classmethod
    def extract_level_from_style(cls, style_name: str) -> int:
        """从样式名提取级别数字"""
        is_heading, level = cls.is_heading_style(style_name)
        return level if is_heading else 0


class TableConverter:
    """表格转换器 - 使用直接XML解析，正确处理rowspan/colspan"""
    
    @staticmethod
    def _tc_info(tc) -> Tuple[int, Optional[str], str]:
        """
        解析单元格信息
        返回: (colspan, vMerge状态, 文本)
        vMerge: 'restart'表示合并起始, 'continue'表示被合并, None表示无合并
        """
        tcPr = tc.find(f'{{{NS}}}tcPr')
        gs = 1
        if tcPr is not None:
            gs_e = tcPr.find(f'{{{NS}}}gridSpan')
            if gs_e is not None:
                gs = int(gs_e.get(f'{{{NS}}}val', 1))
        
        vm = None
        if tcPr is not None:
            vm_e = tcPr.find(f'{{{NS}}}vMerge')
            if vm_e is not None:
                vm = vm_e.get(f'{{{NS}}}val')
                if vm is None:
                    vm = 'continue'  # 无val属性表示continue
        
        # 使用itertext获取纯文本（不是python-docx的cell.text）
        text = ''.join(tc.itertext()).strip()
        return gs, vm, text
    
    @staticmethod
    def _calc_rowspan(cell_data: List[List[Dict]], ri: int, ci: int, n_rows: int,
                      row_grid_starts: List[List[int]]) -> int:
        """
        计算(ri, ci)位置的rowspan
        基于vMerge状态和文本比较
        
        Args:
            cell_data: 单元格数据列表
            ri: 行索引
            ci: tc索引
            n_rows: 总行数
            row_grid_starts: 每行每个tc的grid起始位置列表
        """
        cell = cell_data[ri][ci]
        # 只有restart状态且有文本的单元格才计算rowspan
        if cell['vm'] != 'restart' or not cell['text']:
            return 1
        
        gs = cell['gs']  # gridSpan
        gc_start = row_grid_starts[ri][ci]  # 当前单元格的grid起始位置
        
        rs = 1
        for nr in range(ri + 1, n_rows):
            if nr >= len(cell_data):
                break
            
            # 检查当前单元格占据的每个grid位置
            all_continued = True
            for offset in range(gs):
                gc = gc_start + offset
                # 在后续行中找到占据这个grid位置的单元格
                found_continue = False
                # 遍历后续行的tc，找到占据grid位置gc的单元格
                for next_ci, next_gc in enumerate(row_grid_starts[nr]):
                    if next_ci >= len(cell_data[nr]):
                        continue
                    next_tc = cell_data[nr][next_ci]
                    next_gs = next_tc['gs']
                    if next_gc <= gc < next_gc + next_gs:
                        # 这个tc占据了gc位置
                        # vm='continue' → 延续合并
                        # vm='restart' → 新合并开始，停止
                        # vm=None 且 文本相同 → 可能是无vMerge标记的合并延续
                        if next_tc['vm'] == 'continue':
                            found_continue = True
                        elif next_tc['vm'] is None and next_tc['text'] == cell['text']:
                            found_continue = True
                        else:
                            found_continue = False
                        break
                
                if not found_continue:
                    all_continued = False
                    break
            
            if all_continued:
                rs += 1
            else:
                break
        
        return rs
    
    @staticmethod
    def _infer_colspan(row_cells: List[Dict]) -> List[int]:
        """
        推断每列TC的colspan
        规则：第一个有文本的非合并格，后面连续空TC的数量 → colspan = 空TC数 + 1
        """
        n = len(row_cells)
        colspan = [1] * n
        for i, cell in enumerate(row_cells):
            if cell['text'] and cell['vm'] is None:
                # 第一个有文本的非合并格 → 检查后续连续空格
                empty_after = 0
                for j in range(i + 1, n):
                    if not row_cells[j]['text'] and row_cells[j]['vm'] is None:
                        empty_after += 1
                    else:
                        break
                if empty_after > 0:
                    colspan[i] = empty_after + 1
        return colspan
    
    @staticmethod
    def to_html_table(docx_path: str, table) -> str:
        """
        将表格转换为 HTML 格式（复杂表格）
        
        使用直接XML解析，正确处理:
        - colspan (grid_span > 1)
        - rowspan: 使用vMerge信息
        """
        if not table.rows:
            return ""
        
        # 直接从DOCX解析XML
        try:
            with zipfile.ZipFile(docx_path) as z:
                with z.open('word/document.xml') as f:
                    tree = etree.parse(f)
        except Exception:
            # 解析失败，使用简单的markdown格式
            return TableConverter._to_markdown_table(table)
        
        root = tree.getroot()
        
        # 找到对应的表格（通过行数匹配）
        matched_tbl = None
        tbl_counter = 0
        for tbl in root.findall(f'.//{{{NS}}}tbl'):
            tbl_rows = len(tbl.findall(f'{{{NS}}}tr'))
            if tbl_rows == len(table.rows):
                if tbl_counter == 0:  # 第一个匹配的表格
                    matched_tbl = tbl
                    break
                tbl_counter += 1
        
        if matched_tbl is None:
            return TableConverter._to_markdown_table(table)
        
        tr_list = matched_tbl.findall(f'{{{NS}}}tr')
        n_rows = len(tr_list)
        
        # 获取表格列数（从tblGrid或推算）
        tblGrid = matched_tbl.find(f'{{{NS}}}tblGrid')
        if tblGrid is not None:
            n_cols = len(tblGrid.findall(f'{{{NS}}}gridCol'))
        else:
            n_cols = max(len(tr.findall(f'{{{NS}}}tc')) for tr in tr_list)
        
        # 收集cell_data和每行的grid起始位置
        cell_data: List[List[Dict]] = []
        row_grid_starts: List[List[int]] = []  # 每行每个tc的grid起始位置
        for tr in tr_list:
            row_cells = []
            row_starts = []
            cur_gc = 0
            for tc in tr.findall(f'{{{NS}}}tc'):
                gs, vm, text = TableConverter._tc_info(tc)
                row_cells.append({'text': text, 'gs': gs, 'vm': vm})
                row_starts.append(cur_gc)
                cur_gc += gs
            cell_data.append(row_cells)
            row_grid_starts.append(row_starts)
        
        # remain矩阵：追踪每行每列被rowspan占据的状态
        remain = [[0] * n_cols for _ in range(n_rows)]
        
        lines = ['<table border="1" style="border-collapse:collapse;font-size:12px;">']
        
        for ri, tr in enumerate(tr_list):
            lines.append('  <tr>')
            tcs = tr.findall(f'{{{NS}}}tc')
            tc_grid_start = row_grid_starts[ri]  # 使用预先计算的
            
            for ci, tc in enumerate(tcs):
                gs, vm, text = TableConverter._tc_info(tc)
                gc = tc_grid_start[ci]
                
                # 跳过被rowspan占据的位置
                while gc < n_cols and remain[ri][gc] > 0:
                    remain[ri][gc] -= 1
                    gc += 1
                if gc >= n_cols:
                    break
                
                # continue状态：被垂直合并的单元格，直接跳过
                # （continue单元格的内容来自上方，不需要输出）
                if vm == 'continue':
                    continue
                
                rs = TableConverter._calc_rowspan(cell_data, ri, ci, n_rows, row_grid_starts)
                
                # 预占后续行的位置
                for nr in range(ri + 1, min(ri + rs, n_rows)):
                    remain[nr][gc] = rs - (nr - ri) - 1
                
                # 生成HTML属性
                ca = f' colspan="{gs}"' if gs > 1 else ''
                ra = f' rowspan="{rs}"' if rs > 1 else ''
                tag = 'th' if ri == 0 else 'td'
                safe = text.replace('<', '&lt;').replace('>', '&gt;')
                lines.append(f'    <{tag}{ca}{ra}>{safe}</{tag}>')
            
            lines.append('  </tr>')
        
        lines.append('</table>')
        return '\n'.join(lines)
    
    @staticmethod
    def _to_markdown_table(table) -> str:
        """简单的Markdown表格格式（备用）"""
        lines = []
        for ri, row in enumerate(table.rows):
            md_cols = [cell.text.strip() for cell in row.cells]
            lines.append('| ' + ' | '.join(md_cols) + ' |')
            if ri == 0:
                lines.append('| ' + ' | '.join('---' for _ in row.cells) + ' |')
        return '\n'.join(lines)


class HeaderFooterHandler:
    """页眉页脚处理器"""
    
    def __init__(self, config: ProcessingConfig):
        self.config = config
    
    def extract_header_footer_text(self, doc) -> Set[str]:
        """提取文档页眉页脚中的文本"""
        text_set = set()
        
        if not self.config.remove_page_numbers:
            return text_set
        
        for section in doc.sections:
            # 页眉
            for para in section.header.paragraphs:
                text = para.text.strip()
                if text:
                    text_set.add(text)
            
            # 页脚
            for para in section.footer.paragraphs:
                text = para.text.strip()
                if text:
                    text_set.add(text)
        
        return text_set
    
    def is_page_number(self, text: str) -> bool:
        """判断文本是否为页码"""
        for pattern in self.config.page_number_patterns:
            if re.match(pattern, text, re.IGNORECASE):
                return True
        return False


class ValidationReport:
    """验证报告生成器"""
    
    def __init__(self):
        self.issues: List[str] = []
        self.docx_stats: Dict = {}
        self.md_stats: Dict = {}
    
    def add_issue(self, issue: str):
        self.issues.append(issue)
    
    def set_stats(self, docx_stats: Dict, md_stats: Dict):
        self.docx_stats = docx_stats
        self.md_stats = md_stats
    
    def generate_report(self) -> Dict:
        """生成验证报告"""
        return {
            "docx_stats": self.docx_stats,
            "md_stats": self.md_stats,
            "issues": self.issues,
            "is_valid": len(self.issues) == 0
        }


class NumberingExtractor:
    """
    Word自动编号提取器
    
    从numbering.xml中提取编号格式，用于处理Word自动编号的段落。
    Word的自动编号系统包含：
    - abstractNum: 编号格式抽象定义（格式模板）
    - num: 具体编号实例（应用abstractNum到文档）
    - numPr: 段落中的编号属性引用（numId和ilvl）
    """
    
    def __init__(self, doc):
        """
        初始化编号提取器
        
        Args:
            doc: python-docx Document对象
        """
        self.doc = doc
        self.numbering_part = None
        self.abstract_nums: Dict[str, Dict] = {}  # abstractNumId -> definition
        self.nums: Dict[str, str] = {}  # numId -> abstractNumId
        self.counters: Dict[Tuple[str, int], int] = {}  # (numId, ilvl) -> current count
        
        self._load_numbering()
    
    def _load_numbering(self):
        """加载并解析numbering.xml"""
        # 尝试获取numbering part
        try:
            if hasattr(self.doc.part, 'numbering_part') and self.doc.part.numbering_part:
                self.numbering_part = self.doc.part.numbering_part
                self._parse_numbering_xml()
        except Exception:
            pass  # 没有numbering part是正常的
    
    def _parse_numbering_xml(self):
        """解析numbering XML"""
        if not self.numbering_part:
            return
        
        try:
            import xml.etree.ElementTree as ET
            xml_content = self.numbering_part.blob
            root = ET.fromstring(xml_content)
            
            # 解析abstractNum
            for abstract_num in root.iter():
                tag = abstract_num.tag.split('}')[1] if '}' in abstract_num.tag else abstract_num.tag
                if tag == 'abstractNum':
                    abstract_num_id = abstract_num.get(qn('w:abstractNumId'))
                    if abstract_num_id is None:
                        continue
                    levels = {}
                    
                    for child in abstract_num:
                        child_tag = child.tag.split('}')[1] if '}' in child.tag else child.tag
                        if child_tag == 'lvl':
                            ilvl = child.get(qn('w:ilvl'))
                            if ilvl is None:
                                continue
                            start = None
                            num_fmt = None
                            lvl_text = None
                            
                            for sub_child in child:
                                sub_tag = sub_child.tag.split('}')[1] if '}' in sub_child.tag else sub_child.tag
                                if sub_tag == 'start':
                                    start = sub_child.get(qn('w:val'))
                                elif sub_tag == 'numFmt':
                                    num_fmt = sub_child.get(qn('w:val'))
                                elif sub_tag == 'lvlText':
                                    lvl_text = sub_child.get(qn('w:val'))
                            
                            levels[ilvl] = {
                                'start': int(start) if start else 1,
                                'format': num_fmt or 'decimal',
                                'lvlText': lvl_text or '%1.'
                            }
                    
                    self.abstract_nums[abstract_num_id] = levels
            
            # 解析num（numId -> abstractNumId映射）
            for num_elem in root.iter():
                tag = num_elem.tag.split('}')[1] if '}' in num_elem.tag else num_elem.tag
                if tag == 'num':
                    num_id = num_elem.get(qn('w:numId'))
                    if num_id is None:
                        continue
                    for child in num_elem:
                        child_tag = child.tag.split('}')[1] if '}' in child.tag else child.tag
                        if child_tag == 'abstractNumId':
                            abstract_num_id = child.get(qn('w:val'))
                            if abstract_num_id is not None:
                                self.nums[num_id] = abstract_num_id
        
        except Exception as e:
            pass  # 解析失败时忽略
    
    def get_numbering_info(self, para) -> Optional[Tuple[str, int]]:
        """
        获取段落的编号信息
        
        Args:
            para: 段落对象
            
        Returns:
            (numId, ilvl) 或 None
        """
        p = para._element
        pPr = p.find(qn('w:pPr'))
        if pPr is None:
            return None
        
        numPr = pPr.find(qn('w:numPr'))
        if numPr is None:
            return None
        
        numId_elem = numPr.find(qn('w:numId'))
        ilvl_elem = numPr.find(qn('w:ilvl'))
        
        if numId_elem is None:
            return None
        
        numId = numId_elem.get(qn('w:val'))
        ilvl = ilvl_elem.get(qn('w:val')) if ilvl_elem is not None else '0'
        
        if numId == '0':  # numId=0表示无编号
            return None
        
        return (numId, int(ilvl) if ilvl else 0)
    
    def get_number_text(self, numId: str, ilvl: int) -> str:
        """
        获取编号文本
        
        Args:
            numId: 编号ID
            ilvl: 编号级别
            
        Returns:
            编号文本（如"1."、"1.1"、"a)"等）
        """
        # 获取abstractNumId
        abstract_num_id = self.nums.get(numId)
        if not abstract_num_id:
            return ""
        
        # 获取级别定义
        levels = self.abstract_nums.get(abstract_num_id, {})
        level_def = levels.get(str(ilvl), {})
        
        if not level_def:
            return ""
        
        # 更新计数器
        counter_key = (numId, ilvl)
        if counter_key not in self.counters:
            self.counters[counter_key] = level_def.get('start', 1)
        else:
            self.counters[counter_key] += 1
        
        current_num = self.counters[counter_key]
        
        # 格式化编号
        lvlText = level_def.get('lvlText', '%1.')
        num_fmt = level_def.get('format', 'decimal')
        
        # 替换占位符
        result = lvlText
        for i in range(ilvl + 1):
            placeholder = f'%{i + 1}'
            if placeholder in result:
                # 获取该级别的计数
                level_counter_key = (numId, i)
                level_num = self.counters.get(level_counter_key, 1)
                
                # 根据格式生成编号
                if num_fmt == 'decimal':
                    num_str = str(level_num)
                elif num_fmt == 'lowerLetter':
                    num_str = self._number_to_letter(level_num, lowercase=True)
                elif num_fmt == 'upperLetter':
                    num_str = self._number_to_letter(level_num, lowercase=False)
                elif num_fmt == 'lowerRoman':
                    num_str = self._number_to_roman(level_num, lowercase=True)
                elif num_fmt == 'upperRoman':
                    num_str = self._number_to_roman(level_num, lowercase=False)
                elif num_fmt == 'chineseCountingThousand':
                    num_str = self._number_to_chinese(level_num)
                else:
                    num_str = str(level_num)
                
                result = result.replace(placeholder, num_str)
        
        return result
    
    def _number_to_letter(self, num: int, lowercase: bool = True) -> str:
        """数字转字母 (1=a/1=A, 27=aa/27=AA)"""
        result = ""
        while num > 0:
            num -= 1
            result = chr(97 if lowercase else 65 + num % 26) + result
            num //= 26
        return result
    
    def _number_to_roman(self, num: int, lowercase: bool = True) -> str:
        """数字转罗马数字"""
        val = [1000, 900, 500, 400, 100, 90, 50, 40, 10, 9, 5, 4, 1]
        syb = ['M', 'CM', 'D', 'CD', 'C', 'XC', 'L', 'XL', 'X', 'IX', 'V', 'IV', 'I']
        roman = ''
        for i, v in enumerate(val):
            while num >= v:
                roman += syb[i]
                num -= v
        return roman.lower() if lowercase else roman
    
    def _number_to_chinese(self, num: int) -> str:
        """数字转中文数字"""
        chinese_nums = ['零', '一', '二', '三', '四', '五', '六', '七', '八', '九']
        chinese_units = ['', '十', '百', '千', '万']
        
        if num < 10:
            return chinese_nums[num]
        elif num < 20:
            return '十' + (chinese_nums[num - 10] if num > 10 else '')
        elif num < 100:
            return chinese_nums[num // 10] + '十' + (chinese_nums[num % 10] if num % 10 else '')
        else:
            return str(num)


class ImageExtractor:
    """
    图片提取器 - 从 DOCX 中提取图片并转换为 Base64
    
    DOCX 文件是 ZIP 压缩包，图片存储在 word/media 目录
    每个图片有唯一的 relId，需要通过 document.xml 中的关系查找
    """
    
    # 图片格式到 MIME 类型的映射
    IMAGE_FORMATS = {
        '.png': 'image/png',
        '.jpg': 'image/jpeg',
        '.jpeg': 'image/jpeg',
        '.gif': 'image/gif',
        '.bmp': 'image/bmp',
        '.webp': 'image/webp',
    }
    
    def __init__(self, doc_path: str):
        """
        初始化图片提取器
        
        Args:
            doc_path: DOCX 文件路径
        """
        self.doc_path = doc_path
        self.doc_dir = os.path.dirname(doc_path)
        self._image_cache: Dict[str, str] = {}  # relId -> base64 data URI
        self._raw_cache: Dict[str, bytes] = {}   # relId -> raw image bytes
        self._ext_cache: Dict[str, str] = {}     # relId -> file extension (.png, .jpg)
        self._filename_map: Dict[str, str] = {}  # relId -> original filename
        self._load_images()
    
    def _load_images(self):
        """从 DOCX 中加载所有图片"""
        import zipfile
        import base64
        
        try:
            with zipfile.ZipFile(self.doc_path, 'r') as zf:
                # 读取 word/_rels/document.xml.rels 获取关系
                rels_content = None
                try:
                    rels_content = zf.read('word/_rels/document.xml.rels')
                except KeyError:
                    try:
                        rels_content = zf.read('word/_rels/_rels.document.xml.rels')
                    except KeyError:
                        pass
                
                # 解析关系文件，建立 relId -> 图片路径 的映射
                rel_to_media = {}
                if rels_content:
                    import xml.etree.ElementTree as ET
                    root = ET.fromstring(rels_content)
                    for rel in root.findall('.//{*}Relationship'):
                        rel_id = rel.get('Id')
                        target = rel.get('Target')
                        rel_type = rel.get('Type', '')
                        
                        if 'image' in rel_type.lower() and target:
                            if not target.startswith('media/'):
                                target = 'media/' + target
                            rel_to_media[rel_id] = target
                
                # 提取所有图片
                for filename in zf.namelist():
                    if filename.startswith('word/media/'):
                        image_data = zf.read(filename)
                        base64_data = base64.b64encode(image_data).decode('utf-8')
                        
                        ext = os.path.splitext(filename)[1].lower()
                        mime_type = self.IMAGE_FORMATS.get(ext, 'image/png')
                        data_uri = f"data:{mime_type};base64,{base64_data}"
                        base_name = os.path.basename(filename)
                        
                        # 通过关系找到对应的 relId
                        matched_relid = None
                        for rel_id, media_path in rel_to_media.items():
                            if media_path in filename or filename.endswith(media_path.replace('media/', '')):
                                self._image_cache[rel_id] = data_uri
                                self._raw_cache[rel_id] = image_data
                                self._ext_cache[rel_id] = ext
                                self._filename_map[rel_id] = base_name
                                matched_relid = rel_id
                        
                        # 没找到关系，用序号作为 key
                        if matched_relid is None:
                            key = f"img_{len(self._image_cache)}"
                            self._image_cache[key] = data_uri
                            self._raw_cache[key] = image_data
                            self._ext_cache[key] = ext
                            self._filename_map[key] = base_name
        
        except Exception:
            pass
    
    def get_image_by_relid(self, rel_id: str) -> Optional[str]:
        """通过关系ID获取图片的 Base64 数据（data URI 格式）"""
        return self._image_cache.get(rel_id)
    
    def export_to_file(self, rel_id: str, output_dir: str, counter: int) -> Optional[str]:
        """
        导出图片到文件，返回相对路径
        
        Args:
            rel_id: 关系ID
            output_dir: 输出目录（相对于 doc_dir）
            counter: 图片序号
        
        Returns:
            相对路径如 "images/img_001.png"，或 None
        """
        raw = self._raw_cache.get(rel_id)
        if not raw:
            return None
        
        ext = self._ext_cache.get(rel_id, '.png')
        img_dir = os.path.join(self.doc_dir, output_dir)
        os.makedirs(img_dir, exist_ok=True)
        
        img_name = f"img_{counter:03d}{ext}"
        img_path = os.path.join(img_dir, img_name)
        
        with open(img_path, 'wb') as f:
            f.write(raw)
        
        return f"{output_dir}/{img_name}"
    
    def get_markdown_ref(self, rel_id: str, base64_mode: bool = True,
                         output_dir: str = "images", counter: int = 1) -> str:
        """
        生成 Markdown 图片引用
        
        Args:
            rel_id: 关系ID
            base64_mode: True=内嵌Base64, False=导出文件
            output_dir: 文件导出目录
            counter: 图片序号
        
        Returns:
            Markdown 图片引用字符串
        """
        if base64_mode:
            data_uri = self._image_cache.get(rel_id)
            if data_uri:
                return f"![image]({data_uri})"
            return ""
        else:
            path = self.export_to_file(rel_id, output_dir, counter)
            if path:
                return f"![image]({path})"
            return ""
    
    def get_all_images(self) -> Dict[str, str]:
        """获取所有图片"""
        return self._image_cache.copy()
    
    @staticmethod
    def extract_from_paragraph(paragraph, image_extractor: 'ImageExtractor') -> List[str]:
        """
        从段落中提取图片，返回 Markdown 图片引用列表
        
        Args:
            paragraph: python-docx Paragraph 对象
            image_extractor: ImageExtractor 实例
        
        Returns:
            [markdown_image_ref, ...] 图片引用列表
        """
        images = []
        
        for run in paragraph.runs:
            if hasattr(run, '_element'):
                for drawing in run._element.findall('.//{*}drawing'):
                    for inline in drawing.findall('.//{*}inline'):
                        for blip in inline.findall('.//{*}blip'):
                            embed = blip.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                            if embed:
                                base64_data = image_extractor.get_image_by_relid(embed)
                                if base64_data:
                                    images.append(embed)
        
        return images


class DocxToMarkdownConverter:
    """DOCX 转 Markdown 主转换器"""
    
    def __init__(self, config: Optional[ProcessingConfig] = None):
        self.config = config or ProcessingConfig()
        self.header_footer_handler = HeaderFooterHandler(self.config)
        self.validator = ValidationReport()
        self.numbering_extractor: Optional[NumberingExtractor] = None
        self.image_extractor: Optional[ImageExtractor] = None
        self._image_counter: int = 0
        
        # 跟踪最大标题级别
        self.max_heading_level: int = 1
        self.found_heading_levels: Set[int] = set()
    
    def reset_state(self):
        """重置转换器状态"""
        self.max_heading_level = 1
        self.found_heading_levels = set()
        self.validator = ValidationReport()
        self.numbering_extractor = None
        self.image_extractor = None
        self._image_counter = 0
    
    def convert(self, doc_path: str, output_path: Optional[str] = None) -> Tuple[str, Dict]:
        """
        转换 DOCX 文件为 Markdown
        
        Returns:
            (markdown_content, metadata)
        """
        self.reset_state()
        
        # 导入检查
        try:
            from docx import Document
        except ImportError:
            return "", {"error": "python-docx not installed. Run: pip install python-docx"}
        
        # 加载文档
        try:
            doc = Document(doc_path)
        except Exception as e:
            return "", {"error": f"Failed to load document: {str(e)}"}
        
        # 初始化编号提取器
        self.numbering_extractor = NumberingExtractor(doc)
        
        # 初始化图片提取器
        if self.config.image_enabled:
            self.image_extractor = ImageExtractor(doc_path)
        
        # 提取页眉页脚文本
        header_footer_text = set()
        if self.config.remove_page_numbers:
            header_footer_text = self.header_footer_handler.extract_header_footer_text(doc)
        
        # 收集文档统计信息
        docx_headings = self._count_docx_headings(doc)
        docx_paragraphs = self._count_docx_paragraphs(doc)
        docx_tables = len(doc.tables)
        
        self.validator.docx_stats = {
            "headings": docx_headings,
            "paragraphs": docx_paragraphs,
            "tables": docx_tables
        }
        
        # 处理段落和表格（按文档顺序）
        md_lines = []
        
        # 构建表格元素到表格对象的映射
        table_element_to_obj = {tbl._tbl: tbl for tbl in doc.tables}
        
        # 获取文档正文元素
        body = doc.element.body
        
        # 收集所有段落对象（用于启发式判断）
        all_paragraphs = list(doc.paragraphs)
        para_elements = [p._p for p in all_paragraphs]
        
        # 遍历文档顺序的元素
        last_para_idx = -1
        for child in body:
            tag = child.tag.split('}')[1] if '}' in child.tag else child.tag
            
            if tag == 'p':  # 段落
                # 找到这个段落元素的索引
                try:
                    para_idx = para_elements.index(child)
                except ValueError:
                    continue
                
                para = all_paragraphs[para_idx]
                
                # 获取下一个段落文本（用于启发式判断）
                next_text = ""
                if para_idx + 1 < len(all_paragraphs):
                    next_text = all_paragraphs[para_idx + 1].text.strip()
                
                result = self._process_paragraph(para, header_footer_text, next_text)
                if result is None:
                    continue
                
                para_type, content = result
                
                # 提取图片
                image_md_lines = []
                if self.image_extractor and self.config.image_enabled:
                    img_relids = ImageExtractor.extract_from_paragraph(para, self.image_extractor)
                    for relid in img_relids:
                        self._image_counter += 1
                        md_ref = self.image_extractor.get_markdown_ref(
                            relid,
                            base64_mode=self.config.image_to_base64,
                            output_dir="images",
                            counter=self._image_counter
                        )
                        if md_ref:
                            image_md_lines.append(md_ref)
                
                if para_type == "heading":
                    md_lines.append(content)
                    md_lines.append('')
                    md_lines.extend(image_md_lines)
                
                elif para_type == "list":
                    md_lines.append(content)
                    md_lines.extend(image_md_lines)
                
                elif para_type == "subheading":
                    md_lines.append(f"#### {content}")
                    md_lines.append('')
                    md_lines.extend(image_md_lines)
                
                elif para_type == "paragraph":
                    md_lines.append(content)
                    md_lines.extend(image_md_lines)
                
                elif para_type == "empty":
                    md_lines.extend(image_md_lines)
                elif para_type == "separator":
                    md_lines.extend(image_md_lines)
                
                last_para_idx = para_idx
            
            elif tag == 'tbl' and self.config.enable_table:  # 表格
                # 找到对应的表格对象
                table_obj = table_element_to_obj.get(child)
                if table_obj and table_obj.rows:
                    table_md = TableConverter.to_html_table(doc_path, table_obj)
                    
                    if table_md:
                        md_lines.append(table_md)
                        md_lines.append('')
        
        # 合并为最终 Markdown
        markdown_content = '\n\n'.join(md_lines)
        
        # 修复中文间距
        if self.config.fix_chinese_spacing:
            markdown_content = TextPostProcessor.fix_chinese_spacing(markdown_content)
        
        # 移除多余的空行（清理图片等留下的空行）
        markdown_content = TextPostProcessor.remove_empty_lines(markdown_content)
        
        # 生成 Markdown 统计（正确计数以 # 开头的行）
        md_lines = markdown_content.split('\n')
        md_headings = sum(1 for line in md_lines if line.strip().startswith('#'))
        md_tables = sum(1 for line in md_lines if line.strip().startswith('|') and '---' not in line)
        
        self.validator.md_stats = {
            "headings": md_headings,
            "tables": md_tables
        }
        
        # 一致性检查
        self._validate_conversion(docx_headings, md_headings, docx_tables, md_tables)
        
        # 生成元数据
        metadata = {
            "source": doc_path,
            "format": "docx",
            "chars": len(markdown_content),
            "max_heading_level": self.max_heading_level,
            "found_heading_levels": sorted(self.found_heading_levels),
            "docx_stats": self.validator.docx_stats,
            "md_stats": self.validator.md_stats,
            "validation": self.validator.generate_report()
        }
        
        # 输出到文件
        if output_path:
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(markdown_content)
        
        return markdown_content, metadata
    
    def _process_paragraph(self, para, header_footer_text: Set[str], next_text: str = ""):
        """处理单个段落"""
        text = para.text.strip()
        style_name = para.style.name if para.style and para.style.name else ""
        
        # 跳过空段落（但记录为分隔标记）
        if not text:
            return ("empty", "")
        
        # 跳过页眉页脚中的内容
        if self.config.remove_page_numbers:
            if text in header_footer_text:
                return None
            if self.header_footer_handler.is_page_number(text):
                return None
        
        # 检查是否有Word自动编号
        number_prefix = ""
        if self.numbering_extractor:
            num_info = self.numbering_extractor.get_numbering_info(para)
            if num_info:
                numId, ilvl = num_info
                number_prefix = self.numbering_extractor.get_number_text(numId, ilvl) + " "
        
        # 检测标题（基于样式）
        is_heading, heading_level = HeadingDetector.is_heading_style(style_name)
        
        if is_heading and heading_level > 0:
            # 更新最大标题级别
            self.found_heading_levels.add(heading_level)
            self.max_heading_level = max(self.max_heading_level, heading_level)
            
            # 标题格式化（添加Word自动编号前缀，保留文本中的编号）
            heading_text = f"{number_prefix}{text}" if number_prefix else text
            heading_mark = f"{'#' * heading_level} {heading_text}"
            return ("heading", heading_mark)
        
        # 检测列表
        if self._is_list_paragraph(para):
            # 列表项如果有编号，使用编号而不是默认的"-"
            if number_prefix:
                return ("list", f"{number_prefix.strip()} {text}")
            return ("list", f"- {text}")
        
        # 检测手动编号列表（如"1）xxx"、"2、xxx"、"（3）xxx"）
        if self._is_manual_numbered_list(text):
            return ("list", text)
        
        # 检测目录条目（如"1概述"、"1.1项目综述"、"2.1.3物探方法试验"）
        if self._is_toc_entry(text):
            return ("list", f"- {text}")
        
        # 启发式小标题检测
        if HeadingDetector.is_likely_subheading(text, next_text):
            return ("subheading", text)
        
        # 检测是否是过渡语（后面应该分隔）
        if self._is_transition_text(text):
            return ("separator", "")
        
        # 普通正文段落
        # 检查是否需要首行缩进（2个空格）
        content = text
        pf = para.paragraph_format
        if pf.first_line_indent is not None and pf.first_line_indent > 0:
            content = "  " + text  # 添加2个空格首行缩进
        
        # 如果有编号，添加到段落开头
        if number_prefix:
            content = number_prefix + content
        
        return ("paragraph", content)
    
    def _is_transition_text(self, text: str) -> bool:
        """
        判断是否是过渡语（后面应该有分隔）
        例如："主要包括："、"具体如下："、"本项目建设内容如下：" 等
        
        特征：短文本 + 冒号结尾 + 后面跟着列表项
        """
        if not text or len(text) > 50:
            return False
        
        # 过渡语模式（更精确的匹配）
        transition_patterns = [
            r'如下：$',  # xxx如下：
            r'包括：$',  # xxx包括：
            r'如下$',    # xxx如下
            r'主要内容如下：$',
            r'建设内容如下：$',
            r'如下$',
        ]
        
        for pattern in transition_patterns:
            if re.search(pattern, text):
                return True
        
        return False
    
    def _is_list_paragraph(self, para) -> bool:
        """判断是否为列表段落"""
        style_name = para.style.name if para.style and para.style.name else ""
        if not style_name:
            return False
        
        style_lower = style_name.lower()
        list_indicators = ['list', 'bullet', 'number', '列表', '编号']
        return any(ind in style_lower for ind in list_indicators)
    
    def _is_toc_entry(self, text: str) -> bool:
        """判断是否为目录条目（如"1概述"、"1.1项目综述"、"2.1.3物探方法试验"）"""
        if not text or len(text) > 100:
            return False
        
        # 数字开头 + 中文内容：如 "1概述"、"1.1项目综述"、"2.1.3物探方法试验"
        import re
        toc_pattern = re.compile(r'^\d+(\.\d+)*\s*[\u4e00-\u9fff]')
        if toc_pattern.match(text):
            return True
        
        return False
    
    def _is_manual_numbered_list(self, text: str) -> bool:
        """判断是否为手动编号列表项（如"1）xxx"、"2、xxx"、"（3）xxx"、"1. xxx"）"""
        if not text or len(text) < 3:
            return False
        
        import re
        # 匹配开头的手动编号：数字 + ）、) 、. 、、 等
        patterns = [
            r'^\d+[）\)]',          # 1） 2） 3) 
            r'^\d+[、]',            # 1、 2、 3、
            r'^[（\(]\d+[）\)]',    # （1） (2) （3）
            r'^[①②③④⑤⑥⑦⑧⑨⑩]',  # ① ② ③
            r'^[一二三四五六七八九十]+[、）]',  # 一、 二、 三）
        ]
        for pattern in patterns:
            if re.match(pattern, text):
                return True
        return False
    
    def _process_tables(self, doc, doc_path: str, md_lines: List[str]) -> List[str]:
        """处理文档中的表格"""
        for table in doc.tables:
            if not table.rows:
                continue
            
            table_md = TableConverter.to_html_table(doc_path, table)
            
            if table_md:
                if md_lines and md_lines[-1]:
                    md_lines.append('')
                md_lines.append(table_md)
                md_lines.append('')
        
        return md_lines
    
    def _count_docx_headings(self, doc) -> int:
        """统计 DOCX 中的标题数量"""
        count = 0
        for para in doc.paragraphs:
            if para.style and para.style.name:
                is_heading, level = HeadingDetector.is_heading_style(para.style.name)
                if is_heading:
                    count += 1
        return count
    
    def _count_docx_paragraphs(self, doc) -> int:
        """统计 DOCX 中的段落数量"""
        return sum(1 for p in doc.paragraphs if p.text.strip())
    
    def _validate_conversion(self, docx_headings: int, md_headings: int, 
                             docx_tables: int, md_tables: int):
        """验证转换一致性"""
        # 标题数量检查
        if docx_headings != md_headings:
            self.validator.add_issue(
                f"标题数量不一致: DOCX={docx_headings}, MD={md_headings}"
            )
        
        # 表格数量检查
        if docx_tables != md_tables and md_tables == 0:
            self.validator.add_issue(
                f"表格可能未转换: DOCX={docx_tables}"
            )


def convert_docx_to_markdown(doc_path: str, output_path: Optional[str] = None,
                             config: Optional[ProcessingConfig] = None) -> Tuple[str, Dict]:
    """便捷转换函数"""
    converter = DocxToMarkdownConverter(config)
    return converter.convert(doc_path, output_path)


def main():
    parser = argparse.ArgumentParser(
        description='Enhanced DOCX to Markdown Converter - 通用工具',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
示例:
  python enhanced_convert.py input.docx -o output.md
  python enhanced_convert.py input.docx --json
  python enhanced_convert.py input.docx --merge-paragraphs
  
支持所有 DOCX 文件的通用转换，保留标题级别、段落一致性、表格格式。
        """
    )
    parser.add_argument('input', help='输入 DOCX 文件路径')
    parser.add_argument('-o', '--output', help='输出 Markdown 文件路径')
    parser.add_argument('--json', action='store_true', help='输出元数据为 JSON')
    parser.add_argument('--no-table', action='store_true', help='禁用表格转换')
    parser.add_argument('--no-chinese-fix', action='store_true', help='禁用中文间距修复')
    parser.add_argument('--keep-page-numbers', action='store_true', help='保留页码')
    
    args = parser.parse_args()
    
    config = ProcessingConfig(
        enable_table=not args.no_table,
        fix_chinese_spacing=not args.no_chinese_fix,
        remove_page_numbers=not args.keep_page_numbers,
    )
    
    content, metadata = convert_docx_to_markdown(args.input, args.output, config)
    
    if "error" in metadata:
        print(f"错误: {metadata['error']}", file=sys.stderr)
        sys.exit(1)
    
    if not args.output:
        print(content)
    
    if args.json:
        print(json.dumps(metadata, indent=2, ensure_ascii=False), file=sys.stderr)
    else:
        print(f"\n转换完成:", file=sys.stderr)
        print(f"  标题级别: {metadata.get('max_heading_level', 'N/A')}", file=sys.stderr)
        print(f"  发现级别: {metadata.get('found_heading_levels', [])}", file=sys.stderr)
        print(f"  验证结果: {'通过' if metadata['validation']['is_valid'] else '有问题'}", file=sys.stderr)
        if metadata['validation']['issues']:
            for issue in metadata['validation']['issues']:
                print(f"    - {issue}", file=sys.stderr)


if __name__ == "__main__":
    main()
